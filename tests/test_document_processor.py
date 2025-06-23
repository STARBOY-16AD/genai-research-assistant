import pytest
import os
from pathlib import Path
from src.backend.core.document_processor import DocumentProcessor
from src.backend.core.config import settings
from docx import Document

@pytest.fixture
def doc_processor(tmp_path):
    """Create a DocumentProcessor instance with a temporary vector DB path"""
    # Override VECTOR_DB_PATH to use temporary directory
    settings.VECTOR_DB_PATH = str(tmp_path / "vector_db")
    return DocumentProcessor()

@pytest.fixture
def sample_pdf(tmp_path):
    """Create a sample PDF file for testing"""
    pdf_path = tmp_path / "test.pdf"
    with open(pdf_path, "wb") as f:
        f.write(b"%PDF-1.4\n1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R >>\nendobj\n4 0 obj\n<< /Length 44 >>\nstream\nBT /F1 12 Tf 100 700 Td (This is a test document for GenAI.) Tj ET\nendstream\nendobj\ntrailer\n<< /Root 1 0 R >>\n%%EOF")
    return pdf_path

@pytest.fixture
def sample_txt(tmp_path):
    """Create a sample TXT file for testing"""
    txt_path = tmp_path / "test.txt"
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write("This is a test document.\nIt contains multiple lines.\nFor testing purposes.")
    return txt_path

@pytest.fixture
def sample_docx(tmp_path):
    """Create a sample DOCX file for testing"""
    docx_path = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("This is a test document.")
    doc.add_paragraph("It has multiple paragraphs.")
    doc.save(docx_path)
    return docx_path

@pytest.mark.asyncio
async def test_extract_text_from_pdf(doc_processor, sample_pdf):
    """Test text extraction from PDF"""
    text = doc_processor.extract_text_from_file(str(sample_pdf))
    assert "This is a test document for GenAI" in text
    assert len(text.strip()) > 0

@pytest.mark.asyncio
async def test_extract_text_from_txt(doc_processor, sample_txt):
    """Test text extraction from TXT"""
    text = doc_processor.extract_text_from_file(str(sample_txt))
    assert "This is a test document" in text
    assert "multiple lines" in text
    assert len(text.strip()) > 0

@pytest.mark.asyncio
async def test_extract_text_from_docx(doc_processor, sample_docx):
    """Test text extraction from DOCX"""
    text = doc_processor.extract_text_from_file(str(sample_docx))
    assert "This is a test document" in text
    assert "multiple paragraphs" in text
    assert len(text.strip()) > 0

@pytest.mark.asyncio
async def test_extract_text_unsupported_format(doc_processor, tmp_path):
    """Test extraction with unsupported file format"""
    unsupported_file = tmp_path / "test.jpg"
    with open(unsupported_file, "wb") as f:
        f.write(b"fake image data")
    
    with pytest.raises(ValueError, match="Unsupported file format: .jpg"):
        doc_processor.extract_text_from_file(str(unsupported_file))

@pytest.mark.asyncio
async def test_chunk_document(doc_processor):
    """Test document chunking"""
    sample_text = "This is sentence one. This is sentence two. This is sentence three. " * 10
    doc_id = "test_doc"
    chunks = doc_processor.chunk_document(sample_text, doc_id)
    
    assert len(chunks) > 0
    assert all(chunk["doc_id"] == doc_id for chunk in chunks)
    assert all("id" in chunk and "text" in chunk and "metadata" in chunk for chunk in chunks)
    assert all(len(chunk["text"]) <= settings.CHUNK_SIZE for chunk in chunks)
    assert chunks[0]["chunk_index"] == 0
    assert chunks[-1]["chunk_index"] == len(chunks) - 1

@pytest.mark.asyncio
async def test_store_document_embeddings(doc_processor, tmp_path):
    """Test storing document embeddings in ChromaDB"""
    sample_text = "This is a test document for embedding."
    doc_id = "test_doc"
    chunks = doc_processor.chunk_document(sample_text, doc_id)
    doc_processor.store_document_embeddings(chunks, doc_id)
    
    collection_name = f"doc_{doc_id}"
    collection = doc_processor.chroma_client.get_collection(collection_name)
    assert collection.count() == len(chunks)
    
    # Verify stored documents
    stored_docs = collection.get(include=["documents", "metadatas"])
    assert len(stored_docs["documents"]) == len(chunks)
    assert stored_docs["documents"][0] == chunks[0]["text"]

@pytest.mark.asyncio
async def test_search_relevant_chunks(doc_processor, tmp_path):
    """Test searching relevant chunks"""
    sample_text = "This is a test document about AI. It discusses machine learning."
    doc_id = "test_doc"
    chunks = doc_processor.chunk_document(sample_text, doc_id)
    doc_processor.store_document_embeddings(chunks, doc_id)
    
    query = "What is machine learning?"
    results = doc_processor.search_relevant_chunks(query, doc_id, top_k=2)
    
    assert len(results) > 0
    assert all("text" in result and "id" in result and "distance" in result for result in results)
    assert "machine learning" in results[0]["text"].lower()

@pytest.mark.asyncio
async def test_get_document_hash(doc_processor, sample_txt):
    """Test document hash generation"""
    hash1 = doc_processor.get_document_hash(str(sample_txt))
    hash2 = doc_processor.get_document_hash(str(sample_txt))
    assert hash1 == hash2  # Same file, same hash
    assert len(hash1) == 32  # MD5 hash length

@pytest.mark.asyncio
async def test_process_uploaded_document(doc_processor, sample_txt):
    """Test full document processing pipeline"""
    doc_id, text, chunks = doc_processor.process_uploaded_document(str(sample_txt))
    
    assert doc_id
    assert len(text) > 0
    assert len(chunks) > 0
    assert text.startswith("This is a test document")
    assert all(chunk["doc_id"] == doc_id for chunk in chunks)
    
    # Verify embeddings are stored
    collection = doc_processor.chroma_client.get_collection(f"doc_{doc_id}")
    assert collection.count() == len(chunks)